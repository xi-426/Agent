from pathlib import Path

from docx import Document


DOCX_PATH = Path("docs/知屿个人项目网站完整素材包.docx")
OUTPUT_PATH = DOCX_PATH


REPLACEMENTS = {
    "图2  RAG在线链路：8个候选、0.3496门控、纯向量重排、最终6个上下文":
        "图2  RAG在线链路：pgvector按余弦距离召回Top 8、0.3496门控、无自定义重排",
    "图2  RAG查询链路：8个候选、0.3496门控、90/10重排、最终6条上下文":
        "图2  RAG查询链路：pgvector按余弦距离召回Top 8、0.3496门控、纯向量排序",
    "固定切片后搜索candidateK、finalK、maxDistance和向量/文字权重。":
        "固定切片后搜索Top-K和maxDistance；当前不使用自定义重排。",
    "固定切片后，搜索candidateK、finalK、maxDistance和向量/词面权重。":
        "固定切片后，搜索Top-K和maxDistance；当前不使用自定义重排。",
    "为什么会同时出现91.67%和93.75%？两个指标看的目标不同：Hit@6只看12条有答案题是否召回人工相关来源；门控准确率看全部16条题是否正确决定回答或拒答。":
        "为什么会同时出现91.67%和93.75%？两个指标看的目标不同：Hit@8只看12条有答案题是否在前8条召回人工相关来源；门控准确率看全部16条题是否正确决定回答或拒答。",
    "运行到什么状态再截图：48条标注全部校准，页面显示48、91.7%、93.8%、8→6、阈值0.3496和召回曲线。":
        "运行到什么状态再截图：48条标注全部完成，页面显示48、Hit@8 91.7%、门控93.8%、Top 8、阈值0.3496和召回曲线。",
    "量化结果：12份资料、48条标注、Hit@6 91.7%、门控93.8%，带小样本说明。":
        "量化结果：12份资料、48条标注、Hit@8 91.7%、门控93.8%，带小样本说明。",
    "必须说明HNSW已创建但当前数据规模小，不能声称性能提升；Hit@6不是答案准确率。":
        "必须说明HNSW已创建但当前数据规模小，不能声称性能提升；Hit@8不是答案准确率。",
    "召回、门控、轻量重排、来源返回与拒答":
        "纯向量Top-K、距离门控、来源返回与拒答",
    "召回、门控、轻量重排、来源、两层拒答":
        "纯向量Top-K、距离门控、来源、两层拒答",
    "有答案测试题的相关来源进入最终6个上下文":
        "有答案测试题的相关来源进入前8条检索结果",
    "candidateK / finalK": "Top-K",
    "90%向量 + 10%二字组": "无自定义重排",
    "Hit@6 = 11/12，不是答案准确率": "Hit@8 = 11/12，不是答案准确率",
    "Hit@6": "Hit@8",
    "8→6": "Top 8",
    "8 → 6": "Top 8",
    "90/10": "纯向量",
}


def iter_paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def replace_paragraph_text(paragraph):
    original = paragraph.text
    updated = original
    for old, new in REPLACEMENTS.items():
        updated = updated.replace(old, new)
    if updated == original:
        return False

    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(updated)
    return True


def main():
    document = Document(DOCX_PATH)
    changed = sum(replace_paragraph_text(p) for p in iter_paragraphs(document))
    document.save(OUTPUT_PATH)
    print(f"updated paragraphs: {changed}")


if __name__ == "__main__":
    main()
